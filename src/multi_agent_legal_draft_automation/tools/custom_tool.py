from crewai.tools import BaseTool
from typing import Type, Dict, Any
from pydantic import BaseModel, Field
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re


class TemplateDocToolInput(BaseModel):
    """Input schema for TemplateDocTool."""
    template_path: str = Field(..., description="Path to the Word document template (.docx)")
    output_path: str = Field(..., description="Path where the filled template will be saved")
    replacements: Dict[str, Any] = Field(..., description="Dictionary of placeholder keys and their replacement values")
    placeholder_format: str = Field(default="{{placeholder}}", description="Format of placeholders in the template. Use 'placeholder' to represent the variable name.")


class TemplateDocTool(BaseTool):
    name: str = "Word Template Processor"
    description: str = (
        "A specialized tool for processing Word document templates and replacing placeholders. "
        "This tool preserves ALL formatting and content from the template, "
        "and ONLY replaces the placeholders with the provided values. "
        "It handles tables, paragraphs, headers, and footers, maintaining all formatting including bold and italic text. "
        "Use this tool when you need to fill in a template while maintaining its exact structure and formatting."
    )
    args_schema: Type[BaseModel] = TemplateDocToolInput

    def _run(self, template_path: str, output_path: str, replacements: Dict[str, Any], placeholder_format: str = "{{placeholder}}") -> str:
        try:
            # Ensure the template exists and is a .docx file
            if not os.path.exists(template_path):
                return f"Error: Template not found at path {template_path}"
            if not template_path.endswith('.docx'):
                return f"Error: Template must be a Microsoft Word document (.docx)"

            # Read the template document
            doc = Document(template_path)
            
            # Create placeholder pattern - we'll use {{key}} format
            pattern = re.compile(r"{{(\w+)}}")
            
            # Process paragraphs
            for paragraph in doc.paragraphs:
                self._process_paragraph(paragraph, pattern, replacements)
            
            # Process tables with enhanced table handling
            for table in doc.tables:
                # Process table cells
                for row in table.rows:
                    for cell in row.cells:
                        # Process each paragraph in the cell
                        for paragraph in cell.paragraphs:
                            self._process_paragraph(paragraph, pattern, replacements)
            
            # Process headers and footers
            for section in doc.sections:
                for header in section.header.paragraphs:
                    self._process_paragraph(header, pattern, replacements)
                for footer in section.footer.paragraphs:
                    self._process_paragraph(footer, pattern, replacements)
            
            # Save the document
            doc.save(output_path)
            return f"Successfully processed template and saved document at {output_path}"

        except Exception as e:
            return f"Error processing template: {str(e)}"
    
    def _process_paragraph(self, paragraph, pattern, replacements):
        # Check if paragraph contains any placeholders
        if not pattern.search(paragraph.text):
            return
        
        # Store the paragraph's style and alignment
        style = paragraph.style
        alignment = paragraph.alignment
        
        # Store information about each run and its formatting
        runs_info = []
        for run in paragraph.runs:
            runs_info.append({
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_size': run.font.size,
                'font_name': run.font.name,
                'color': run.font.color.rgb if run.font.color else None
            })
        
        # Get the original text
        original_text = paragraph.text
        
        # Replace all placeholders in the full text
        new_text = original_text
        for key, value in replacements.items():
            placeholder = "{{" + key + "}}"
            new_text = new_text.replace(placeholder, str(value))
        
        # If no changes were made, return
        if new_text == original_text:
            return
        
        # Clear the paragraph
        paragraph.clear()
        
        # If there were no runs with formatting, just add the new text
        if not runs_info:
            paragraph.add_run(new_text)
            paragraph.style = style
            paragraph.alignment = alignment
            return
        
        # Try to preserve formatting by finding original text segments in the new text
        current_pos = 0
        for run_info in runs_info:
            original_segment = run_info['text']
            if not original_segment:  # Skip empty runs
                continue
                
            # Find where this segment appears in the new text
            pos = new_text.find(original_segment, current_pos)
            
            # If the segment is found, add text before it and then the segment with its formatting
            if pos >= 0:
                # Add any text before this segment (likely replaced placeholder text)
                if pos > current_pos:
                    paragraph.add_run(new_text[current_pos:pos])
                
                # Add the segment with its original formatting
                run = paragraph.add_run(original_segment)
                self._apply_formatting(run, run_info)
                
                # Update position
                current_pos = pos + len(original_segment)
            else:
                # If segment not found, it might have been part of a placeholder
                # We'll handle remaining text at the end
                continue
        
        # Add any remaining text
        if current_pos < len(new_text):
            paragraph.add_run(new_text[current_pos:])
        
        # Restore style and alignment
        paragraph.style = style
        paragraph.alignment = alignment
    
    def _apply_formatting(self, run, formatting):
        """Apply formatting to a run."""
        run.bold = formatting['bold']
        run.italic = formatting['italic']
        
        # Apply additional formatting if available
        if formatting['underline']:
            run.underline = formatting['underline']
        
        if formatting['font_size']:
            run.font.size = formatting['font_size']
            
        if formatting['font_name']:
            run.font.name = formatting['font_name']
            
        if formatting['color']:
            run.font.color.rgb = formatting['color']
